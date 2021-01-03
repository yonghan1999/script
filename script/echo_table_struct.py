import MySQLdb
from docx import Document
from docx.shared import Inches
from docx.shared import  Pt


print("please enter host(e.g localhost)\n")
host = input("your host:") or "localhost"
print("please enter user(e.g root)\n")
user = input("user name:") or "root"
print("please enter password(e.g 123456)\n")
password = input("password:") or "123456"
print("please enter database name(e.g DBName)\n")
dbname = input("database name:") or "exam"
print("please enter database charset(e.g utf8)\n")
charset = input("charset:") or "utf8"

document = Document()
try:
    db = MySQLdb.connect(host, user, password, dbname, charset=charset)
    cursor = db.cursor()

    cursor.execute("SHOW TABLES")
    rows = cursor.fetchall()
    tables = []
    
    for row in rows:
        tables.append(row)
    for table in tables:
        data = []
        cursor.execute("DESC " + table[0])
        rows = cursor.fetchall()
        for row in rows:
            data.append(row)
            
        paragraph = document.add_paragraph()
        paragraph.add_run(table[0]).font.size=Pt(21)

        
        form = document.add_table(rows=1, cols=6, style='Table Grid')
        hdr_cells = form.rows[0].cells
        hdr_cells[0].text = "字段名"
        hdr_cells[1].text = "长度"
        hdr_cells[2].text = "允许空值"
        hdr_cells[3].text = "键"
        hdr_cells[4].text = "虚拟"
        hdr_cells[5].text = "默认"
        for line in data:
            row_cells = form.add_row().cells
            row_cells[0].text = line[0]
            row_cells[1].text = line[1]
            row_cells[2].text = line[2]
            row_cells[3].text = line[3]
            row_cells[4].text = line[4] or 'None'
            row_cells[5].text = line[5]
    document.save('table.docx')
    print("The file is named tables.docx in the same directory with the script")
      
    
        
except MySQLdb.Error as err:
    print("Mysql Error: {0}".format(err))
    

